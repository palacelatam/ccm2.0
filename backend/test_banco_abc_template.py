"""
Test script specifically for the Banco ABC template
"""
import asyncio
import sys
import os
from datetime import datetime

# Add src to path
sys.path.append('src')

from services.settlement_instruction_service import settlement_instruction_service


async def test_banco_abc_template():
    """Test the Banco ABC template with sample data"""
    
    print("=" * 60)
    print("Testing Template Carta Instrucción Banco ABC")
    print("=" * 60)
    
    # Sample trade data with all possible fields
    trade_data = {
        # Standard trade fields
        'TradeNumber': 'ABC-2025-001',
        'client_name': 'XYZ Corporation',
        'CounterpartyName': 'Banco ABC',
        'Product': 'FX_FORWARD',
        'Currency1': 'USD',
        'Currency2': 'CLP',
        'QuantityCurrency1': 1000000,
        'QuantityCurrency2': 950000000,
        'Price': 950.00,
        'Direction': 'SELL',
        'TradeDate': '2025-01-09',
        'MaturityDate': '2025-02-09',
        
        # Additional fields that might be in the template
        'trade_id': 'ABC-2025-001',
        'client_id': 'xyz-corp',
        'client_rut': '76.123.456-7',
        'client_address': 'Av. Providencia 1234, Santiago',
        'bank_rut': '97.004.000-5',
        'settlement_date': '2025-02-09',
        'spot_rate': 945.50,
        'forward_points': 4.50,
        'trader_name': 'Juan Pérez',
        'sales_person': 'María González',
        'confirmation_number': 'CONF-2025-001',
    }
    
    # Sample settlement data with bank account info
    settlement_data = {
        'account_name': 'XYZ Corporation - USD Account',
        'account_number': '1234567890',
        'bank_name': 'Banco de Chile',
        'swift_code': 'BCHICLRM',
        'cutoff_time': '15:00 Santiago Time',
        'special_instructions': 'Please notify treasury@xyzcorp.cl upon receipt of funds.',
        
        # Additional settlement fields
        'intermediary_bank': 'JPMorgan Chase',
        'intermediary_swift': 'CHASUS33',
        'beneficiary_name': 'XYZ Corporation',
        'beneficiary_address': 'Av. Providencia 1234, Santiago, Chile',
        'reference': 'FX Forward Settlement - Trade ABC-2025-001',
    }
    
    print("\nTrade Data:")
    print("-" * 40)
    for key, value in trade_data.items():
        print(f"  {key}: {value}")
    
    print("\nSettlement Data:")
    print("-" * 40)
    for key, value in settlement_data.items():
        print(f"  {key}: {value}")
    
    print("\nGenerating settlement instruction...")
    print("-" * 40)
    
    try:
        # Generate the document
        result = await settlement_instruction_service.generate_settlement_instruction(
            trade_data=trade_data,
            template_name='Template Carta Instrucción Banco ABC.docx',
            settlement_data=settlement_data
        )
        
        if result['success']:
            print("SUCCESS: Document generated!")
            print(f"  Path: {result['document_path']}")
            print(f"  Template: {result['template_used']}")
            print(f"  Variables populated: {result['variables_populated']}")
            print(f"  Generated at: {result['generated_at']}")
            
            # Show the filename for easy access
            filename = os.path.basename(result['document_path'])
            print(f"\nOpen the document to verify:")
            print(f"  backend/templates/generated/{filename}")
            
        else:
            print("ERROR: Document generation failed!")
            print(f"  Error: {result['error']}")
            
    except Exception as e:
        print(f"ERROR: Exception occurred: {e}")
        import traceback
        traceback.print_exc()
    
    print("\n" + "=" * 60)


async def check_template_exists():
    """Check if the template file exists and is readable"""
    template_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        'templates',
        'Template Carta Instrucción Banco ABC.docx'
    )
    
    print(f"\nChecking template file:")
    print(f"  Path: {template_path}")
    print(f"  Exists: {os.path.exists(template_path)}")
    
    if os.path.exists(template_path):
        file_size = os.path.getsize(template_path)
        print(f"  Size: {file_size:,} bytes")
        
        # Try to open it with python-docx to verify it's valid
        try:
            from docx import Document
            doc = Document(template_path)
            print(f"  Paragraphs: {len(doc.paragraphs)}")
            print(f"  Tables: {len(doc.tables)}")
            
            # Show first few paragraphs to see content
            print("\n  First 5 paragraphs:")
            for i, para in enumerate(doc.paragraphs[:5], 1):
                text = para.text.strip()
                if text:
                    # Truncate long lines
                    if len(text) > 60:
                        text = text[:60] + "..."
                    print(f"    {i}. {text}")
                    
        except Exception as e:
            print(f"  ERROR reading document: {e}")
    else:
        print("  ERROR: Template file not found!")
        print("\n  Please ensure 'Template Carta Instrucción Banco ABC.docx'")
        print("  is placed in the backend/templates/ directory")


async def main():
    """Main function"""
    # First check if template exists
    await check_template_exists()
    
    # Then test the template
    await test_banco_abc_template()


if __name__ == "__main__":
    asyncio.run(main())
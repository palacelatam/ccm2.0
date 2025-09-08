"""
Settlement Instruction Service for generating populated settlement instruction documents
"""

from typing import Optional, Dict, Any, List
from datetime import datetime
import os
import logging
from docx import Document
from docx.shared import Inches
import tempfile
import uuid
from services.storage_service import StorageService
from config.firebase_config import get_cmek_firestore_client
from utils.bank_utils import get_bank_display_name

logger = logging.getLogger(__name__)


class SettlementInstructionService:
    """Service for generating populated settlement instruction documents"""
    
    def __init__(self):
        """Initialize the service"""
        # Get templates directory path
        self.templates_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
            'templates'
        )
        
        # Ensure templates directory exists
        os.makedirs(self.templates_dir, exist_ok=True)
        
        # Initialize CMEK Firestore client (same as rest of application)
        self.db = get_cmek_firestore_client()
        
        # Initialize storage service
        self.storage_service = StorageService()
        
        logger.info(f"Settlement Instruction Service initialized with templates dir: {self.templates_dir}")
    
    async def _resolve_counterparty_to_bank_id(self, client_id: str, counterparty_name: str) -> str:
        """
        Resolve counterparty name to bank ID using client-specific mappings.
        
        Args:
            client_id: Client ID for mapping lookup
            counterparty_name: Original counterparty name from trade
            
        Returns:
            Resolved bank ID or original counterparty name if no mapping found
        """
        if not counterparty_name or not client_id:
            return counterparty_name or ""
        
        try:
            # Check client's counterparty mappings (case-insensitive)
            mappings_ref = self.db.collection('clients').document(client_id).collection('counterpartyMappings')
            
            # First try exact case match
            mappings_query = mappings_ref.where('counterpartyName', '==', counterparty_name).limit(1)
            mappings = list(mappings_query.stream())
            
            # If no exact match, try case-insensitive search by getting all mappings
            if not mappings:
                all_mappings = list(mappings_ref.stream())
                for mapping_doc in all_mappings:
                    mapping_data = mapping_doc.to_dict()
                    stored_name = mapping_data.get('counterpartyName', '')
                    if stored_name.lower() == counterparty_name.lower():
                        mappings = [mapping_doc]
                        break
            
            if mappings:
                mapping_data = mappings[0].to_dict()
                bank_id = mapping_data.get('bankId')
                if bank_id:
                    logger.info(f"✅ Resolved counterparty '{counterparty_name}' → '{bank_id}' via client mapping")
                    return bank_id
            
            # Fallback: return original counterparty name
            logger.info(f"⚠️ No mapping found for counterparty '{counterparty_name}', using as-is")
            return counterparty_name
            
        except Exception as e:
            logger.error(f"❌ Error resolving counterparty '{counterparty_name}': {e}")
            return counterparty_name
    
    def create_standard_template(self) -> str:
        """
        Create a standard settlement instruction template for testing
        
        Returns:
            Path to created template file
        """
        template_path = os.path.join(self.templates_dir, 'standard_settlement_instruction.docx')
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading('SETTLEMENT INSTRUCTIONS', 0)
        title.alignment = 1  # Center alignment
        
        # Add document content with placeholders
        doc.add_paragraph('')  # Blank line
        
        # Client information section
        doc.add_heading('Client Information', level=1)
        p1 = doc.add_paragraph()
        p1.add_run('Client Name: ').bold = True
        p1.add_run('{client_name}')
        
        # Trade information section
        doc.add_heading('Trade Details', level=1)
        
        trade_info = [
            ('Trade Number:', '{trade_number}'),
            ('Trade Date:', '{trade_date}'),
            ('Value Date:', '{value_date}'),
            ('Counterparty:', '{counterparty_name}'),
            ('Product:', '{product_type}'),
            ('Currency Pair:', '{currency_1}/{currency_2}'),
            ('Amount:', '{amount_currency_1} {currency_1}'),
            ('Counter Amount:', '{amount_currency_2} {currency_2}'),
            ('Exchange Rate:', '{exchange_rate}'),
            ('Direction:', '{direction}')
        ]
        
        for label, placeholder in trade_info:
            p = doc.add_paragraph()
            p.add_run(label + ' ').bold = True
            p.add_run(placeholder)
        
        # Settlement information section
        doc.add_heading('Settlement Instructions', level=1)
        
        settlement_info = [
            ('Account Name:', '{account_name}'),
            ('Account Number:', '{account_number}'),
            ('Bank Name:', '{bank_name}'),
            ('SWIFT Code:', '{swift_code}'),
            ('Cutoff Time:', '{cutoff_time}'),
        ]
        
        for label, placeholder in settlement_info:
            p = doc.add_paragraph()
            p.add_run(label + ' ').bold = True
            p.add_run(placeholder)
        
        # Special instructions section
        doc.add_heading('Special Instructions', level=1)
        doc.add_paragraph('{special_instructions}')
        
        # Footer
        doc.add_paragraph('')
        footer_p = doc.add_paragraph('Generated automatically by Client Confirmation Manager 2.0')
        footer_p.alignment = 1  # Center alignment
        
        # Save template
        doc.save(template_path)
        logger.info(f"Standard template created at: {template_path}")
        
        return template_path
    
    async def find_matching_template(
        self,
        bank_id: str,
        trade_data: Dict[str, Any],
        client_segment_id: Optional[str] = None
    ) -> Optional[Dict[str, Any]]:
        """
        Find the best matching settlement instruction template from the database
        
        Args:
            bank_id: Bank ID to search templates for
            trade_data: Trade information for matching
            client_segment_id: Optional client segment ID for more specific matching
            
        Returns:
            Template document data or None if no match found
        """
        try:
            # Get settlement type from trade data
            settlement_type = trade_data.get('SettlementType', trade_data.get('settlement_type', ''))
            product = trade_data.get('Product', trade_data.get('product', ''))
            
            logger.info(f"🔍 Template search criteria:")
            logger.info(f"   Bank ID: '{bank_id}'")
            logger.info(f"   Settlement type: '{settlement_type}'")
            logger.info(f"   Product: '{product}'")
            logger.info(f"   Client segment ID: '{client_segment_id}'")
            
            # Query settlement instruction letters from database
            logger.info(f"🔍 Querying database: banks/{bank_id}/settlementInstructionLetters")
            letters_ref = self.db.collection('banks').document(bank_id).collection('settlementInstructionLetters')
            
            # First, let's see if the bank document exists
            bank_doc = self.db.collection('banks').document(bank_id).get()
            if not bank_doc.exists:
                logger.error(f"❌ Bank document 'banks/{bank_id}' does not exist in database!")
                return None
            
            logger.info(f"✅ Bank document exists: {bank_doc.to_dict()}")
            
            # Check if there are any settlement instruction letters at all
            all_letters_query = letters_ref.limit(10)  # Just get first 10 to check
            all_letters = list(all_letters_query.stream())
            
            logger.info(f"🔍 Total settlement instruction letters in banks/{bank_id}/settlementInstructionLetters: {len(all_letters)}")
            
            if len(all_letters) == 0:
                logger.error(f"❌ No settlement instruction letters found for bank '{bank_id}' at all!")
                return None
            
            # Show all available templates for debugging
            logger.info(f"📋 Available templates for bank '{bank_id}':")
            for i, letter in enumerate(all_letters, 1):
                data = letter.to_dict()
                logger.info(f"   {i}. ID: {letter.id}")
                logger.info(f"      Rule name: {data.get('rule_name', 'N/A')}")
                logger.info(f"      Settlement type: {data.get('settlement_type', 'N/A')}")
                logger.info(f"      Product: {data.get('product', 'N/A')}")
                logger.info(f"      Client segment: {data.get('client_segment_id', 'default')}")
                logger.info(f"      Active: {data.get('active', 'N/A')}")
                logger.info(f"      Priority: {data.get('priority', 'N/A')}")
                logger.info(f"      Storage path: {data.get('document_storage_path', 'N/A')}")
            
            # Start with active templates only
            logger.info(f"🔍 Filtering by active=True...")
            query = letters_ref.where('active', '==', True)
            active_templates = list(query.stream())
            logger.info(f"   Found {len(active_templates)} active templates")
            
            # Filter by settlement type if available
            if settlement_type:
                logger.info(f"🔍 Filtering by settlement_type='{settlement_type}'...")
                query = query.where('settlement_type', '==', settlement_type)
                settlement_filtered = list(query.stream())
                logger.info(f"   Found {len(settlement_filtered)} templates matching settlement_type")
            else:
                settlement_filtered = active_templates
                logger.info(f"⚠️ No settlement type provided, using all active templates")
            
            # Get all matching templates
            templates = query.stream()
            
            # Convert to list and sort by priority and specificity
            template_list = []
            for template in templates:
                template_data = template.to_dict()
                template_data['id'] = template.id
                
                logger.info(f"🧮 Scoring template '{template_data.get('rule_name')}':")
                
                # Calculate match score (higher = better match)
                score = 0
                
                # Exact settlement type match
                template_settlement_type = template_data.get('settlement_type', '')
                if template_settlement_type == settlement_type:
                    score += 100
                    logger.info(f"   Settlement type match: +100 ('{template_settlement_type}' == '{settlement_type}')")
                else:
                    logger.info(f"   Settlement type mismatch: +0 ('{template_settlement_type}' != '{settlement_type}')")
                    
                # Product match
                template_product = template_data.get('product', '')
                if template_product and product:
                    if template_product.upper() in product.upper():
                        score += 50
                        logger.info(f"   Product partial match: +50 ('{template_product}' in '{product}')")
                    elif product.upper() in template_product.upper():
                        score += 30
                        logger.info(f"   Product partial match: +30 ('{product}' in '{template_product}')")
                    else:
                        logger.info(f"   Product mismatch: +0 ('{template_product}' vs '{product}')")
                else:
                    logger.info(f"   Product not specified: +0")
                    
                # Client segment match (most specific)
                template_client_segment = template_data.get('client_segment_id')
                if client_segment_id and template_client_segment == client_segment_id:
                    score += 200
                    logger.info(f"   Client segment exact match: +200 ('{template_client_segment}' == '{client_segment_id}')")
                elif not template_client_segment:
                    score += 10
                    logger.info(f"   Default template (no client segment): +10")
                else:
                    logger.info(f"   Client segment mismatch: +0 ('{template_client_segment}' vs '{client_segment_id}')")
                    
                # Priority from database (lower number = higher priority)
                priority = template_data.get('priority', 999)
                priority_score = (1000 - priority)
                score += priority_score
                logger.info(f"   Priority bonus: +{priority_score} (priority {priority})")
                
                template_data['match_score'] = score
                logger.info(f"   Total score: {score}")
                template_list.append(template_data)
                
            # Sort by match score (descending)
            template_list.sort(key=lambda x: x['match_score'], reverse=True)
            
            logger.info(f"🏆 Final template ranking:")
            for i, template in enumerate(template_list, 1):
                logger.info(f"   {i}. {template.get('rule_name')} (score: {template['match_score']})")
            
            if template_list:
                best_match = template_list[0]
                logger.info(f"✅ Selected best template: {best_match.get('rule_name')} (score: {best_match['match_score']})")
                return best_match
            else:
                logger.error(f"❌ No templates passed the filtering criteria!")
                return None
                
        except Exception as e:
            logger.error(f"💥 Error finding matching template: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return None
    
    async def download_template_from_storage(self, template_data: Dict[str, Any]) -> Optional[str]:
        """
        Download template document from cloud storage to local temp file
        
        Args:
            template_data: Template document data from database
            
        Returns:
            Path to downloaded template file or None if failed
        """
        try:
            storage_path = template_data.get('document_storage_path')
            if not storage_path:
                logger.error(f"No storage path found for template {template_data.get('id')}")
                return None
                
            # Create temporary file
            temp_dir = tempfile.gettempdir()
            temp_filename = f"template_{template_data.get('id', 'unknown')}_{uuid.uuid4().hex}.docx"
            temp_path = os.path.join(temp_dir, temp_filename)
            
            # Download from storage
            download_result = await self.storage_service.download_document(
                storage_path=storage_path,
                local_path=temp_path
            )
            
            if download_result.get('success'):
                logger.info(f"Template downloaded to: {temp_path}")
                return temp_path
            else:
                logger.error(f"Failed to download template: {download_result.get('error')}")
                return None
                
        except Exception as e:
            logger.error(f"Error downloading template from storage: {e}")
            return None
    
    def populate_template(
        self, 
        template_path: str, 
        trade_data: Dict[str, Any],
        settlement_data: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        Populate a template with trade and settlement data
        
        Args:
            template_path: Path to the template document
            trade_data: Dictionary containing trade information
            settlement_data: Optional dictionary containing settlement account information
            
        Returns:
            Path to the populated document
        """
        try:
            # Load template
            doc = Document(template_path)
            
            # Prepare variable mapping (pass template name for language detection)
            template_name = os.path.basename(template_path)
            variables = self._prepare_variables(trade_data, settlement_data, template_name)
            
            # Replace placeholders in all paragraphs
            for paragraph in doc.paragraphs:
                self._replace_placeholders_in_paragraph(paragraph, variables)
            
            # Replace placeholders in tables (if any)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholders_in_paragraph(paragraph, variables)
            
            # Generate unique filename for populated document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            trade_id = trade_data.get('trade_number', 'unknown')
            filename = f"settlement_instruction_{trade_id}_{timestamp}.docx"
            output_path = os.path.join(self.templates_dir, 'generated', filename)
            
            # Ensure generated directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Save populated document
            doc.save(output_path)
            
            logger.info(f"Settlement instruction generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error populating template: {e}")
            raise
    
    def _prepare_variables(
        self, 
        trade_data: Dict[str, Any], 
        settlement_data: Optional[Dict[str, Any]] = None,
        template_name: Optional[str] = None
    ) -> Dict[str, str]:
        """
        Prepare variable mapping from trade and settlement data
        
        Args:
            trade_data: Trade information
            settlement_data: Settlement account information
            template_name: Template filename to detect language
            
        Returns:
            Dictionary mapping placeholder names to values
        """
        variables = {}
        
        # Detect language from template name (default to Spanish)
        language = 'es'  # Default Spanish
        if template_name:
            template_lower = template_name.lower()
            if 'english' in template_lower or '_en' in template_lower or '-en' in template_lower:
                language = 'en'
            elif 'portuguese' in template_lower or 'portugues' in template_lower or '_pt' in template_lower or '-pt' in template_lower:
                language = 'pt'
        
        # Localization dictionaries
        direction_translations = {
            'en': {'BUY': 'Buy', 'SELL': 'Sell'},
            'es': {'BUY': 'Compra', 'SELL': 'Venta'},
            'pt': {'BUY': 'Compra', 'SELL': 'Venda'}
        }
        
        action_translations = {
            'en': {'credit': 'credit', 'debit': 'debit'},
            'es': {'credit': 'abonar', 'debit': 'cargar'},
            'pt': {'credit': 'creditar', 'debit': 'debitar'}
        }
        
        # Trade data variables
        trade_direction = str(trade_data.get('Direction', trade_data.get('direction', ''))).upper()
        
        # Get localized direction
        localized_direction = direction_translations[language].get(trade_direction, trade_direction)
        
        # Determine action words for Entrega Física based on trade direction
        # If Buy: we receive (credit) currency1 and pay (debit) currency2
        # If Sell: we pay (debit) currency1 and receive (credit) currency2
        if trade_direction == 'BUY':
            action_currency_1 = action_translations[language]['credit']  # receiving currency1
            action_currency_2 = action_translations[language]['debit']   # paying currency2
        else:  # SELL
            action_currency_1 = action_translations[language]['debit']   # paying currency1
            action_currency_2 = action_translations[language]['credit']  # receiving currency2
        
        variables.update({
            'client_name': str(trade_data.get('client_name', 'N/A')),
            'trade_number': str(trade_data.get('TradeNumber', trade_data.get('trade_number', 'N/A'))),
            'trade_date': self._format_date(trade_data.get('TradeDate', trade_data.get('trade_date'))),
            'value_date': self._format_date(trade_data.get('MaturityDate', trade_data.get('value_date'))),
            'counterparty_name': str(trade_data.get('CounterpartyName', trade_data.get('counterparty', 'N/A'))),
            'product_type': str(trade_data.get('Product', trade_data.get('product_type', 'N/A'))),
            'currency_1': str(trade_data.get('Currency1', trade_data.get('currency_1', 'N/A'))),
            'currency_2': str(trade_data.get('Currency2', trade_data.get('currency_2', 'N/A'))),
            'amount_currency_1': self._format_amount(trade_data.get('QuantityCurrency1', trade_data.get('amount_currency_1'))),
            'amount_currency_2': self._format_amount(trade_data.get('QuantityCurrency2', trade_data.get('amount_currency_2'))),
            'price': str(trade_data.get('Price', trade_data.get('exchange_rate', 'N/A'))),
            'settlement_currency': str(trade_data.get('SettlementCurrency', trade_data.get('settlement_currency', 'N/A'))),
            'settlement_type': str(trade_data.get('SettlementType', trade_data.get('settlement_type', 'N/A'))),
            'direction': localized_direction,  # Use localized direction
            'direction_original': str(trade_data.get('Direction', trade_data.get('direction', 'N/A'))),  # Keep original for reference
            'action_currency_1': action_currency_1,
            'action_currency_2': action_currency_2,
            'client_name': str(trade_data.get('client_name', 'N/A')),
            'todays_date': datetime.now().strftime('%d/%m/%Y')
        })
        
        # Settlement data variables (if provided)
        if settlement_data:
            # Check if this is physical delivery (has cargar/abonar accounts) or compensación (single account)
            if 'cargar_account_number' in settlement_data and 'abonar_account_number' in settlement_data:
                # Physical delivery - we have two accounts (cargar and abonar)
                # Add the cargar/abonar structure directly for template variables
                variables.update({
                    # Cargar account (pay/outflow)
                    'cargar_account_name': str(settlement_data.get('cargar_account_name', 'N/A')),
                    'cargar_account_number': str(settlement_data.get('cargar_account_number', 'N/A')),
                    'cargar_bank_name': get_bank_display_name(str(settlement_data.get('cargar_bank_name', 'N/A'))),
                    'cargar_swift_code': str(settlement_data.get('cargar_swift_code', 'N/A')),
                    'cargar_currency': str(settlement_data.get('cargar_currency', 'N/A')),
                    
                    # Abonar account (receive/inflow)
                    'abonar_account_name': str(settlement_data.get('abonar_account_name', 'N/A')),
                    'abonar_account_number': str(settlement_data.get('abonar_account_number', 'N/A')),
                    'abonar_bank_name': get_bank_display_name(str(settlement_data.get('abonar_bank_name', 'N/A'))),
                    'abonar_swift_code': str(settlement_data.get('abonar_swift_code', 'N/A')),
                    'abonar_currency': str(settlement_data.get('abonar_currency', 'N/A')),
                    
                    # Generic fields for basic templates (use abonar account as primary)
                    'account_name': str(settlement_data.get('abonar_account_name', 'N/A')),
                    'account_number': str(settlement_data.get('abonar_account_number', 'N/A')),
                    'bank_name': get_bank_display_name(str(settlement_data.get('abonar_bank_name', 'N/A'))),
                    'swift_code': str(settlement_data.get('abonar_swift_code', 'N/A')),
                })
                
                # Map cargar/abonar to currency1/currency2 based on trade direction and currencies
                currency_1 = str(trade_data.get('Currency1', 'N/A'))
                currency_2 = str(trade_data.get('Currency2', 'N/A'))
                cargar_currency = str(settlement_data.get('cargar_currency', 'N/A'))
                abonar_currency = str(settlement_data.get('abonar_currency', 'N/A'))
                
                # Map accounts to currencies based on which currency they hold
                if cargar_currency == currency_1:
                    # Cargar account holds currency_1
                    variables.update({
                        'account_number_currency_1': str(settlement_data.get('cargar_account_number', 'N/A')),
                        'account_bank_currency_1': get_bank_display_name(str(settlement_data.get('cargar_bank_name', 'N/A'))),
                        'account_name_currency_1': str(settlement_data.get('cargar_account_name', 'N/A')),
                        'swift_code_currency_1': str(settlement_data.get('cargar_swift_code', 'N/A')),
                        'account_number_currency_2': str(settlement_data.get('abonar_account_number', 'N/A')),
                        'account_bank_currency_2': get_bank_display_name(str(settlement_data.get('abonar_bank_name', 'N/A'))),
                        'account_name_currency_2': str(settlement_data.get('abonar_account_name', 'N/A')),
                        'swift_code_currency_2': str(settlement_data.get('abonar_swift_code', 'N/A')),
                    })
                else:
                    # Abonar account holds currency_1
                    variables.update({
                        'account_number_currency_1': str(settlement_data.get('abonar_account_number', 'N/A')),
                        'account_bank_currency_1': get_bank_display_name(str(settlement_data.get('abonar_bank_name', 'N/A'))),
                        'account_name_currency_1': str(settlement_data.get('abonar_account_name', 'N/A')),
                        'swift_code_currency_1': str(settlement_data.get('abonar_swift_code', 'N/A')),
                        'account_number_currency_2': str(settlement_data.get('cargar_account_number', 'N/A')),
                        'account_bank_currency_2': get_bank_display_name(str(settlement_data.get('cargar_bank_name', 'N/A'))),
                        'account_name_currency_2': str(settlement_data.get('cargar_account_name', 'N/A')),
                        'swift_code_currency_2': str(settlement_data.get('cargar_swift_code', 'N/A')),
                    })
                
                # Add settlement data fields
                variables.update({
                    'cutoff_time': str(settlement_data.get('cutoff_time', 'N/A')),
                    'special_instructions': str(settlement_data.get('special_instructions', 'Standard settlement instructions apply.')),
                    'central_bank_trade_code': str(settlement_data.get('central_bank_trade_code', 'N/A'))
                })
            else:
                # Compensación - single account
                # Add single account fields
                variables.update({
                    'account_name': str(settlement_data.get('account_name', 'N/A')),
                    'account_number': str(settlement_data.get('account_number', 'N/A')),
                    'account_bank': get_bank_display_name(str(settlement_data.get('bank_name', 'N/A'))),
                    'swift_code': str(settlement_data.get('swift_code', 'N/A')),
                    'cutoff_time': str(settlement_data.get('cutoff_time', 'N/A')),
                    'special_instructions': str(settlement_data.get('special_instructions', 'Standard settlement instructions apply.')),
                    'central_bank_trade_code': str(settlement_data.get('central_bank_trade_code', 'N/A'))
                })
                
                # Also add cargar/abonar fields using the single account for both (for template compatibility)
                variables.update({
                    # Cargar account (same as main account for Compensación)
                    'cargar_account_name': str(settlement_data.get('account_name', 'N/A')),
                    'cargar_account_number': str(settlement_data.get('account_number', 'N/A')),
                    'cargar_bank_name': get_bank_display_name(str(settlement_data.get('bank_name', 'N/A'))),
                    'cargar_swift_code': str(settlement_data.get('swift_code', 'N/A')),
                    'cargar_currency': str(settlement_data.get('settlement_currency', trade_data.get('SettlementCurrency', trade_data.get('settlement_currency', 'N/A')))),
                    
                    # Abonar account (same as main account for Compensación)
                    'abonar_account_name': str(settlement_data.get('account_name', 'N/A')),
                    'abonar_account_number': str(settlement_data.get('account_number', 'N/A')),
                    'abonar_bank_name': get_bank_display_name(str(settlement_data.get('bank_name', 'N/A'))),
                    'abonar_swift_code': str(settlement_data.get('swift_code', 'N/A')),
                    'abonar_currency': str(settlement_data.get('settlement_currency', trade_data.get('SettlementCurrency', trade_data.get('settlement_currency', 'N/A')))),
                    
                    # Currency-based mapping for backwards compatibility
                    'account_number_currency_1': str(settlement_data.get('account_number', 'N/A')),
                    'account_bank_currency_1': get_bank_display_name(str(settlement_data.get('bank_name', 'N/A'))),
                    'account_name_currency_1': str(settlement_data.get('account_name', 'N/A')),
                    'swift_code_currency_1': str(settlement_data.get('swift_code', 'N/A')),
                    'account_number_currency_2': str(settlement_data.get('account_number', 'N/A')),
                    'account_bank_currency_2': get_bank_display_name(str(settlement_data.get('bank_name', 'N/A'))),
                    'account_name_currency_2': str(settlement_data.get('account_name', 'N/A')),
                    'swift_code_currency_2': str(settlement_data.get('swift_code', 'N/A')),
                })
        else:
            # Default settlement data
            variables.update({
                'account_name': 'N/A',
                'account_number': 'N/A',
                'bank_name': 'N/A',
                'swift_code': 'N/A',
                'account_number_currency_1': 'N/A',
                'account_bank_currency_1': 'N/A',
                'account_number_currency_2': 'N/A',
                'account_bank_currency_2': 'N/A',
                'cutoff_time': 'N/A',
                'special_instructions': 'Standard settlement instructions apply.',
                'central_bank_trade_code': 'N/A',
                # Default cargar/abonar fields
                'cargar_account_name': 'N/A',
                'cargar_account_number': 'N/A',
                'cargar_bank_name': 'N/A',
                'cargar_swift_code': 'N/A',
                'cargar_currency': 'N/A',
                'abonar_account_name': 'N/A',
                'abonar_account_number': 'N/A',
                'abonar_bank_name': 'N/A',
                'abonar_swift_code': 'N/A',
                'abonar_currency': 'N/A'
            })
        
        return variables
    
    def _replace_placeholders_in_paragraph(self, paragraph, variables: Dict[str, str]):
        """Replace placeholders in a paragraph with actual values"""
        full_text = paragraph.text
        
        for placeholder, value in variables.items():
            placeholder_with_braces = '{' + placeholder + '}'
            if placeholder_with_braces in full_text:
                full_text = full_text.replace(placeholder_with_braces, value)
        
        # Only update if text changed
        if full_text != paragraph.text:
            # Try to preserve formatting from the first run
            if paragraph.runs:
                first_run = paragraph.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                bold = first_run.font.bold
                italic = first_run.font.italic
                underline = first_run.font.underline
                
                # Clear and update text
                paragraph.clear()
                new_run = paragraph.add_run(full_text)
                
                # Apply the preserved formatting
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                if bold:
                    new_run.font.bold = bold
                if italic:
                    new_run.font.italic = italic
                if underline:
                    new_run.font.underline = underline
            else:
                # No existing runs, just update
                paragraph.clear()
                paragraph.add_run(full_text)
    
    def _format_date(self, date_value: Any) -> str:
        """Format date value for display"""
        if not date_value:
            return 'N/A'
        
        if isinstance(date_value, str):
            # Try to parse common date formats
            for fmt in ['%Y-%m-%d', '%d-%m-%Y', '%d/%m/%Y', '%m/%d/%Y']:
                try:
                    parsed_date = datetime.strptime(date_value, fmt)
                    return parsed_date.strftime('%d/%m/%Y')
                except ValueError:
                    continue
            return str(date_value)  # Return as-is if parsing fails
        
        if isinstance(date_value, datetime):
            return date_value.strftime('%d/%m/%Y')
        
        return str(date_value)
    
    def _format_amount(self, amount_value: Any) -> str:
        """Format amount value for display"""
        if not amount_value:
            return 'N/A'
        
        try:
            # Try to format as number with thousands separator
            if isinstance(amount_value, (int, float)):
                return f"{amount_value:,.2f}"
            
            # Try to parse string as number
            if isinstance(amount_value, str):
                # Remove common separators and try parsing
                cleaned = amount_value.replace(',', '').replace(' ', '')
                if cleaned.replace('.', '').replace('-', '').isdigit():
                    num_value = float(cleaned)
                    return f"{num_value:,.2f}"
            
            return str(amount_value)
        except (ValueError, TypeError):
            return str(amount_value)
    
    async def generate_settlement_instruction(
        self,
        trade_data: Dict[str, Any],
        bank_id: str,
        client_segment_id: Optional[str] = None,
        settlement_data: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Main method to generate a settlement instruction document
        
        Args:
            trade_data: Trade information dictionary
            bank_id: Bank ID to find appropriate template for
            client_segment_id: Optional client segment ID for more specific template matching
            settlement_data: Optional settlement account information
            
        Returns:
            Dictionary with generation results
        """
        try:
            logger.info(f"🔍 Starting settlement instruction generation:")
            logger.info(f"   Bank ID: {bank_id}")
            logger.info(f"   Client segment ID: {client_segment_id}")
            logger.info(f"   Settlement type: {trade_data.get('SettlementType', trade_data.get('settlement_type', 'N/A'))}")
            logger.info(f"   Product: {trade_data.get('Product', trade_data.get('product', 'N/A'))}")
            
            # Find matching template from database
            template_data = await self.find_matching_template(
                bank_id=bank_id,
                trade_data=trade_data,
                client_segment_id=client_segment_id
            )
            
            if not template_data:
                error_msg = f"❌ No matching template found in database for bank_id='{bank_id}', settlement_type='{trade_data.get('SettlementType', trade_data.get('settlement_type'))}', product='{trade_data.get('Product', trade_data.get('product'))}'"
                logger.error(error_msg)
                return {
                    'success': False,
                    'error': error_msg,
                    'generated_at': datetime.now().isoformat()
                }
            
            logger.info(f"✅ Found matching template: {template_data.get('rule_name')} (ID: {template_data.get('id')}, Score: {template_data.get('match_score')})")
            
            # Download template from cloud storage
            template_path = await self.download_template_from_storage(template_data)
            template_name = template_data.get('rule_name', 'Unknown Template')
            
            if not template_path:
                error_msg = f"❌ Failed to download template '{template_name}' from cloud storage"
                logger.error(error_msg)
                return {
                    'success': False,
                    'error': error_msg,
                    'template_id': template_data.get('id'),
                    'generated_at': datetime.now().isoformat()
                }
            
            logger.info(f"✅ Template downloaded to: {template_path}")
            
            # Populate template
            populated_doc_path = self.populate_template(template_path, trade_data, settlement_data)
            logger.info(f"✅ Document generated: {populated_doc_path}")
            
            # Clean up downloaded template file if it was temporary
            if template_path.startswith(tempfile.gettempdir()):
                try:
                    os.remove(template_path)
                    logger.info(f"🧹 Cleaned up temporary template file: {template_path}")
                except Exception as e:
                    logger.warning(f"⚠️ Could not clean up temporary file {template_path}: {e}")
            
            # Get the variables for debugging
            debug_variables = self._prepare_variables(trade_data, settlement_data)
            
            return {
                'success': True,
                'document_path': populated_doc_path,
                'template_used': template_name,
                'template_id': template_data.get('id'),
                'match_score': template_data.get('match_score'),
                'variables_populated': len(debug_variables),
                'debug_variables': debug_variables,  # Add this for debugging
                'generated_at': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"💥 Failed to generate settlement instruction: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return {
                'success': False,
                'error': str(e),
                'generated_at': datetime.now().isoformat()
            }
    
    async def find_and_prepare_settlement_data(self, trade_data: Dict[str, Any], settlement_rules: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
        """
        Find matching settlement rules and prepare settlement data for a trade.
        SHARED function used by both manual and automatic flows.
        Returns complete settlement data ready for document generation.
        
        Args:
            trade_data: Trade data dictionary
            settlement_rules: List of settlement rules
            
        Returns:
            Complete settlement data dictionary or None
        """
        if not settlement_rules:
            return None
            
        original_counterparty = trade_data.get('CounterpartyName', trade_data.get('BankCounterparty', ''))
        
        # Resolve counterparty name to bank ID using client mappings BEFORE rule matching
        resolved_counterparty = await self._resolve_counterparty_to_bank_id(
            trade_data.get('client_id', ''), 
            original_counterparty
        )
        
        trade_counterparty = resolved_counterparty.lower()
        trade_currency1 = trade_data.get('Currency1', '')
        trade_currency2 = trade_data.get('Currency2', '')
        trade_product = trade_data.get('ProductType', trade_data.get('Product', ''))
        trade_direction = trade_data.get('Direction', '').upper()
        settlement_type = trade_data.get('SettlementType', '')
        settlement_currency = trade_data.get('SettlementCurrency', '')
        
        logger.info(f"🔍 Matching settlement rules - Type: {settlement_type}, Currency: {settlement_currency}")
        logger.info(f"   Counterparty: '{trade_counterparty}'")
        logger.info(f"   Product: '{trade_product}'")
        logger.info(f"   Direction: '{trade_direction}'")
        logger.info(f"   Currency1: '{trade_currency1}'")
        logger.info(f"   Currency2: '{trade_currency2}'")
        logger.info(f"   Available rules: {len(settlement_rules)}")
        
        # Debug: Show all rules
        for i, rule in enumerate(settlement_rules):
            logger.info(f"   Rule {i+1}/{len(settlement_rules)}: name='{rule.get('name', 'unnamed')}', "
                       f"counterparty='{rule.get('counterparty', '')}', "
                       f"product='{rule.get('product', '')}', "
                       f"modalidad='{rule.get('modalidad', '')}', "
                       f"settlementCurrency='{rule.get('settlementCurrency', '')}', "
                       f"active={rule.get('active', True)}")
        
        if settlement_type == "Compensación":
            for rule in settlement_rules:
                rule_counterparty = rule.get('counterparty', '').lower()
                rule_product = rule.get('product', '')
                rule_modalidad = rule.get('modalidad', '')
                rule_settlement_currency = rule.get('settlementCurrency', '')
                rule_active = rule.get('active', True)
                
                if not rule_active:
                    continue
                    
                # Check if all non-empty criteria match
                counterparty_matches = not rule_counterparty or rule_counterparty in trade_counterparty
                product_matches = not rule_product or rule_product.lower() in trade_product.lower()
                # Case-insensitive comparison and handle accents (ó vs o)
                modalidad_matches = not rule_modalidad or rule_modalidad.lower().replace('ó', 'o') == settlement_type.lower().replace('ó', 'o')
                currency_matches = not rule_settlement_currency or rule_settlement_currency == settlement_currency
                
                # Debug each match
                logger.info(f"      Rule '{rule.get('name', 'unnamed')}': "
                           f"counterparty={counterparty_matches} ({rule_counterparty} in {trade_counterparty}), "
                           f"product={product_matches} ({rule_product} in {trade_product}), "
                           f"modalidad={modalidad_matches} ({rule_modalidad} == {settlement_type}), "
                           f"currency={currency_matches} ({rule_settlement_currency} == {settlement_currency})")
                
                if counterparty_matches and product_matches and modalidad_matches and currency_matches:
                    logger.info(f"✅ Matched rule: {rule.get('name')} for Compensación")
                    
                    # Return complete settlement data from the rule
                    return {
                        'matched_rule': rule,
                        # Basic fields
                        'account_name': rule.get('abonarAccountName', 'N/A'),
                        'account_number': rule.get('abonarAccountNumber', 'N/A'),
                        'bank_name': rule.get('abonarBankName', 'N/A'),  # Keep original ID for internal processing
                        'swift_code': rule.get('abonarSwiftCode', 'N/A'),
                        
                        # Abonar fields
                        'abonar_account_name': rule.get('abonarAccountName', 'N/A'),
                        'abonar_account_number': rule.get('abonarAccountNumber', 'N/A'),
                        'abonar_bank_name': rule.get('abonarBankName', 'N/A'),  # Keep original ID for internal processing
                        'abonar_swift_code': rule.get('abonarSwiftCode', 'N/A'),
                        'abonar_currency': rule.get('abonarCurrency', 'N/A'),
                        
                        # Cargar fields (same as abonar for Compensación)
                        'cargar_account_name': rule.get('abonarAccountName', 'N/A'),
                        'cargar_account_number': rule.get('cargarAccountNumber', rule.get('abonarAccountNumber', 'N/A')),
                        'cargar_bank_name': rule.get('cargarBankName', rule.get('abonarBankName', 'N/A')),  # Keep original ID for internal processing
                        'cargar_swift_code': rule.get('cargarSwiftCode', rule.get('abonarSwiftCode', 'N/A')),
                        'cargar_currency': rule.get('cargarCurrency', rule.get('abonarCurrency', 'N/A')),
                        
                        # Other fields
                        'cutoff_time': '15:00 Santiago Time',
                        'special_instructions': rule.get('specialInstructions', 'Standard settlement instructions apply.'),
                        'central_bank_trade_code': rule.get('centralBankTradeCode', 'N/A')
                    }
                    
        elif settlement_type == "Entrega Física":
            # Determine pay/receive currencies based on direction
            if trade_direction == "BUY":
                pay_currency = trade_currency2
                receive_currency = trade_currency1
            else:
                pay_currency = trade_currency1
                receive_currency = trade_currency2
                
            for rule in settlement_rules:
                rule_counterparty = rule.get('counterparty', '').lower()
                rule_product = rule.get('product', '')
                rule_modalidad = rule.get('modalidad', '')
                cargar_currency = rule.get('cargarCurrency', '')
                abonar_currency = rule.get('abonarCurrency', '')
                rule_active = rule.get('active', True)
                
                if not rule_active:
                    continue
                    
                # Check all criteria
                counterparty_matches = not rule_counterparty or rule_counterparty in trade_counterparty
                product_matches = not rule_product or rule_product.lower() in trade_product.lower()
                # Case-insensitive comparison and handle accents (ó vs o)
                modalidad_matches = not rule_modalidad or rule_modalidad.lower().replace('ó', 'o') == settlement_type.lower().replace('ó', 'o')
                currency_matches = cargar_currency == pay_currency and abonar_currency == receive_currency
                
                if counterparty_matches and product_matches and modalidad_matches and currency_matches:
                    logger.info(f"✅ Matched rule: {rule.get('name')} for Entrega Física")
                    
                    # Return complete settlement data from the rule
                    return {
                        'matched_rule': rule,
                        'pay_currency': pay_currency,
                        'receive_currency': receive_currency,
                        
                        # Cargar account fields
                        'cargar_account_name': rule.get('cargarAccountName', 'N/A'),
                        'cargar_account_number': rule.get('cargarAccountNumber', 'N/A'),
                        'cargar_bank_name': rule.get('cargarBankName', 'N/A'),  # Keep original ID for internal processing
                        'cargar_swift_code': rule.get('cargarSwiftCode', 'N/A'),
                        'cargar_currency': rule.get('cargarCurrency', 'N/A'),
                        
                        # Abonar account fields
                        'abonar_account_name': rule.get('abonarAccountName', 'N/A'),
                        'abonar_account_number': rule.get('abonarAccountNumber', 'N/A'),
                        'abonar_bank_name': rule.get('abonarBankName', 'N/A'),  # Keep original ID for internal processing
                        'abonar_swift_code': rule.get('abonarSwiftCode', 'N/A'),
                        'abonar_currency': rule.get('abonarCurrency', 'N/A'),
                        
                        # Basic fields (for backward compatibility)
                        'account_name': rule.get('abonarAccountName', 'N/A'),
                        'account_number': rule.get('abonarAccountNumber', 'N/A'),
                        'bank_name': rule.get('abonarBankName', 'N/A'),  # Keep original ID for internal processing
                        'swift_code': rule.get('abonarSwiftCode', 'N/A'),
                        
                        # Other fields
                        'cutoff_time': '15:00 Santiago Time',
                        'special_instructions': 'Physical delivery settlement - two-way transfer',
                        'central_bank_trade_code': rule.get('centralBankTradeCode', 'N/A')
                    }
        
        return None
    
    async def update_email_with_settlement_path(self, client_id: str, email_id: str, storage_path: str, trade_index: int = 0) -> bool:
        """
        Update email document in Firestore with settlement instruction storage path.
        SHARED function used by both manual and automatic flows.
        
        Args:
            client_id: Client ID
            email_id: Email document ID
            storage_path: Cloud storage path of the settlement instruction
            trade_index: Index of the trade in the Trades array (default 0)
            
        Returns:
            True if update successful, False otherwise
        """
        try:
            db = get_cmek_firestore_client()
            
            # Update the email document with the settlement instruction storage path
            email_doc_ref = db.collection('clients').document(client_id).collection('emails').document(email_id)
            email_doc = email_doc_ref.get()
            
            if email_doc.exists:
                email_data = email_doc.to_dict()
                llm_data = email_data.get('llmExtractedData', {})
                trades = llm_data.get('Trades', [])
                
                # Update the specific trade in the array while preserving array structure
                if trade_index < len(trades) and isinstance(trades, list):
                    trades[trade_index]['settlementInstructionStoragePath'] = storage_path
                    
                    # Update the entire Trades array to preserve its structure
                    email_doc_ref.update({
                        'llmExtractedData.Trades': trades
                    })
                    logger.info(f"✅ Settlement instruction storage path stored in Firestore for email {email_id}, trade {trade_index}")
                    return True
                else:
                    logger.error(f"❌ Invalid trade index {trade_index} or Trades is not an array")
                    return False
            else:
                logger.warning(f"⚠️ Email document {email_id} does not exist - settlement instruction metadata not attached to email")
                return False
                
        except Exception as e:
            logger.error(f"❌ Failed to update email with settlement path: {e}")
            return False


# Global instance
settlement_instruction_service = SettlementInstructionService()